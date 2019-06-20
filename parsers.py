def parse_docx(file, i):
	import docx 
	try:
		doc = docx.Document(file)
		with open('op\\' + str(i) + ".txt", "a", encoding="utf-8") as op:
			for para in doc.paragraphs:
				op.write(para.text + "\n")
			for table in doc.tables:
				for row in table.rows:
				    for cell in row.cells:
				        for paragraph in cell.paragraphs:
				        	op.write(paragraph.text + "\n")
		return True
	except:
		import traceback
		traceback.print_exc()
		with open("err.txt", "a", encoding="utf-8") as op:
			op.write("Failed: " + file + "\n")

def explore(dir):
	import os
	files = []
	for f in os.listdir(dir):
		if(os.path.isdir(dir + "\\" + f) and dir != "op" and dir[0] not in ["!","_"]):
			files = files + explore(dir + "\\" + f)
		elif(os.path.isfile(dir + "\\" + f)):
			print(f)
			files.append(dir + "\\" + f)
	return files

if __name__ == '__main__':
 	import os
 	print(explore(os.getcwd()))